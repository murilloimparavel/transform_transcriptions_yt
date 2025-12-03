"""
Módulo para extrair texto de diferentes fontes: sites, PDFs, documentos Word, arquivos de texto (.txt), Markdown (.md), Excel (.xlsx, .xls) e CSV (.csv).
"""
import os
import re
import csv
import requests
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
from termcolor import colored, cprint
try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False

def extract_text_from_url(url, timeout=30):
    """
    Extrai texto de uma URL (web scraping).
    
    Args:
        url: URL do site
        timeout: Timeout em segundos
    
    Returns:
        str: Texto extraído do site
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        cprint(f"🌐 Acessando URL: {url}", "blue")
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove scripts e styles
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.decompose()
        
        # Extrai texto
        text = soup.get_text()
        
        # Limpa o texto (remove espaços múltiplos, quebras de linha excessivas)
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        if not text or len(text) < 100:
            raise ValueError("Texto extraído muito curto ou vazio")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except requests.exceptions.RequestException as e:
        raise ValueError(f"Erro ao acessar URL: {e}")
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto da URL: {e}")


def extract_text_from_pdf(pdf_path):
    """
    Extrai texto de um arquivo PDF.
    
    Args:
        pdf_path: Caminho do arquivo PDF
    
    Returns:
        str: Texto extraído do PDF
    """
    try:
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Arquivo PDF não encontrado: {pdf_path}")
        
        cprint(f"📄 Lendo PDF: {pdf_path}", "blue")
        
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            cprint(f"   Total de páginas: {total_pages}", "white")
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                text += page_text + "\n"
                
                if page_num % 10 == 0:
                    cprint(f"   Processando página {page_num}/{total_pages}...", "white")
        
        if not text or len(text.strip()) < 100:
            raise ValueError("PDF não contém texto suficiente ou está vazio")
        
        # Limpa o texto
        text = re.sub(r'\s+', ' ', text)
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do PDF: {e}")


def extract_text_from_docx(docx_path):
    """
    Extrai texto de um arquivo Word (.docx).
    
    Args:
        docx_path: Caminho do arquivo Word
    
    Returns:
        str: Texto extraído do documento
    """
    try:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Arquivo Word não encontrado: {docx_path}")
        
        cprint(f"📝 Lendo documento Word: {docx_path}", "blue")
        
        doc = Document(docx_path)
        
        text_parts = []
        
        # Extrai texto dos parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extrai texto das tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        if not text or len(text.strip()) < 100:
            raise ValueError("Documento Word não contém texto suficiente ou está vazio")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do Word: {e}")


def extract_text_from_txt(txt_path):
    """
    Extrai texto de um arquivo de texto (.txt).
    
    Args:
        txt_path: Caminho do arquivo de texto
    
    Returns:
        str: Texto extraído do arquivo
    """
    try:
        if not os.path.exists(txt_path):
            raise FileNotFoundError(f"Arquivo de texto não encontrado: {txt_path}")
        
        cprint(f"📄 Lendo arquivo de texto: {txt_path}", "blue")
        
        # Tenta diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        text = None
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    text = file.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError(f"Não foi possível ler o arquivo com nenhum encoding suportado")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Arquivo de texto está vazio ou contém muito pouco conteúdo")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do arquivo: {e}")


def extract_text_from_markdown(md_path):
    """
    Extrai texto de um arquivo Markdown (.md).
    
    Args:
        md_path: Caminho do arquivo Markdown
    
    Returns:
        str: Texto extraído do arquivo
    """
    try:
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"Arquivo Markdown não encontrado: {md_path}")
        
        cprint(f"📝 Lendo arquivo Markdown: {md_path}", "blue")
        
        # Tenta diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        text = None
        
        for encoding in encodings:
            try:
                with open(md_path, 'r', encoding=encoding) as file:
                    text = file.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError(f"Não foi possível ler o arquivo com nenhum encoding suportado")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Arquivo Markdown está vazio ou contém muito pouco conteúdo")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do arquivo Markdown: {e}")


def extract_text_from_excel(excel_path):
    """
    Extrai texto de um arquivo Excel (.xlsx ou .xls).
    
    Args:
        excel_path: Caminho do arquivo Excel
    
    Returns:
        str: Texto extraído do arquivo
    """
    try:
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Arquivo Excel não encontrado: {excel_path}")
        
        ext = os.path.splitext(excel_path)[1].lower()
        cprint(f"📊 Lendo arquivo Excel: {excel_path}", "blue")
        
        text_parts = []
        
        # Processa .xlsx usando openpyxl
        if ext == '.xlsx':
            if not OPENPYXL_AVAILABLE:
                raise ValueError("Biblioteca openpyxl não está instalada. Execute: pip install openpyxl")
            
            workbook = load_workbook(excel_path, data_only=True)
            total_sheets = len(workbook.sheetnames)
            
            cprint(f"   Total de planilhas: {total_sheets}", "white")
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            # Converte o valor para string e limpa
                            cell_str = str(cell_value).strip()
                            if cell_str:
                                row_text.append(cell_str)
                    if row_text:
                        sheet_text.append(" | ".join(row_text))
                
                if sheet_text:
                    text_parts.append(f"[Planilha: {sheet_name}]")
                    text_parts.extend(sheet_text)
                    text_parts.append("")  # Linha em branco entre planilhas
        
        # Processa .xls usando xlrd
        elif ext == '.xls':
            if not XLRD_AVAILABLE:
                raise ValueError("Biblioteca xlrd não está instalada. Execute: pip install xlrd")
            
            workbook = xlrd.open_workbook(excel_path)
            total_sheets = workbook.nsheets
            
            cprint(f"   Total de planilhas: {total_sheets}", "white")
            
            for sheet_idx in range(total_sheets):
                sheet = workbook.sheet_by_index(sheet_idx)
                sheet_name = sheet.name
                sheet_text = []
                
                for row_idx in range(sheet.nrows):
                    row = sheet.row(row_idx)
                    row_text = []
                    for cell in row:
                        if cell.value:
                            cell_str = str(cell.value).strip()
                            if cell_str:
                                row_text.append(cell_str)
                    if row_text:
                        sheet_text.append(" | ".join(row_text))
                
                if sheet_text:
                    text_parts.append(f"[Planilha: {sheet_name}]")
                    text_parts.extend(sheet_text)
                    text_parts.append("")  # Linha em branco entre planilhas
        else:
            raise ValueError(f"Extensão de arquivo Excel não suportada: {ext}. Use .xlsx ou .xls")
        
        text = "\n".join(text_parts)
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Arquivo Excel não contém dados suficientes ou está vazio")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do Excel: {e}")


def extract_text_from_csv(csv_path):
    """
    Extrai texto de um arquivo CSV (.csv).
    
    Args:
        csv_path: Caminho do arquivo CSV
    
    Returns:
        str: Texto extraído do arquivo
    """
    try:
        if not os.path.exists(csv_path):
            raise FileNotFoundError(f"Arquivo CSV não encontrado: {csv_path}")
        
        cprint(f"📋 Lendo arquivo CSV: {csv_path}", "blue")
        
        # Tenta diferentes encodings e delimitadores
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        delimiters = [',', ';', '\t']
        text_parts = []
        success = False
        
        for encoding in encodings:
            if success:
                break
            for delimiter in delimiters:
                try:
                    with open(csv_path, 'r', encoding=encoding, newline='') as file:
                        # Detecta o delimitador lendo a primeira linha
                        sample = file.read(1024)
                        file.seek(0)
                        
                        # Tenta detectar automaticamente o delimitador
                        sniffer = csv.Sniffer()
                        try:
                            detected_delimiter = sniffer.sniff(sample, delimiters=delimiters).delimiter
                        except:
                            detected_delimiter = delimiter
                        
                        reader = csv.reader(file, delimiter=detected_delimiter)
                        
                        for row in reader:
                            # Filtra células vazias
                            row_text = [cell.strip() for cell in row if cell.strip()]
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    
                    success = True
                    break
                except (UnicodeDecodeError, csv.Error):
                    continue
                except Exception:
                    continue
        
        if not text_parts:
            raise ValueError(f"Não foi possível ler o arquivo CSV com nenhum encoding ou delimitador suportado")
        
        text = "\n".join(text_parts)
        
        if not text or len(text.strip()) < 10:
            raise ValueError("Arquivo CSV está vazio ou contém muito pouco conteúdo")
        
        cprint(f"✅ Texto extraído: {len(text)} caracteres", "green")
        return text
        
    except FileNotFoundError:
        raise
    except Exception as e:
        raise ValueError(f"Erro ao extrair texto do CSV: {e}")


def extract_text_from_source(source_path):
    """
    Detecta automaticamente o tipo de fonte e extrai o texto.
    
    Args:
        source_path: URL, caminho de PDF, Word, arquivo de texto (.txt), Markdown (.md), Excel (.xlsx, .xls) ou CSV (.csv)
    
    Returns:
        tuple: (texto_extraído, tipo_fonte)
    """
    # Verifica se é URL
    if source_path.startswith(('http://', 'https://')):
        text = extract_text_from_url(source_path)
        return text, "url"
    
    # Verifica extensão do arquivo
    ext = os.path.splitext(source_path)[1].lower()
    
    if ext == '.pdf':
        text = extract_text_from_pdf(source_path)
        return text, "pdf"
    elif ext in ['.docx', '.doc']:
        if ext == '.doc':
            raise ValueError("Arquivos .doc não são suportados. Use .docx")
        text = extract_text_from_docx(source_path)
        return text, "docx"
    elif ext == '.txt':
        text = extract_text_from_txt(source_path)
        return text, "txt"
    elif ext == '.md':
        text = extract_text_from_markdown(source_path)
        return text, "md"
    elif ext in ['.xlsx', '.xls']:
        text = extract_text_from_excel(source_path)
        return text, "excel"
    elif ext == '.csv':
        text = extract_text_from_csv(source_path)
        return text, "csv"
    else:
        raise ValueError(f"Formato não suportado: {ext}. Use URL, PDF (.pdf), Word (.docx), Texto (.txt), Markdown (.md), Excel (.xlsx, .xls) ou CSV (.csv)")

